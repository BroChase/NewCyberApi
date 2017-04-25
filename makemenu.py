import os
from tkinter import filedialog, Menu, messagebox
import datetime
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet
import requests
from docx import Document
from docx.shared import Inches



class MenuMaker2000():
    def __init__(self, tkapp):
        self.tkapp = tkapp
        self.saveFilename = None
        #self.exportMenu("docx")

    def createMenu(self):
        topmenu = Menu(self.tkapp)

        filemenu = Menu(topmenu, tearoff=0)
        filemenu.add_command(label="Open", command=self.openMenu)
        filemenu.add_command(label="Save", command=lambda: self.saveMenu("Save"))
        filemenu.add_command(label="Save As", command=lambda: self.saveMenu("Save As"))
        filemenu.add_separator()

        exportmenu = Menu(topmenu, tearoff=0)
        exportmenu.add_command(label="pdf", command=lambda: self.exportMenu("pdf"))
        exportmenu.add_command(label="docx", command=lambda: self.exportMenu("docx"))

        filemenu.add_cascade(label="Export", menu=exportmenu)
        filemenu.add_separator()

        filemenu.add_command(label="Quit", command=self.tkapp.quit)
        topmenu.add_cascade(label="File", menu=filemenu)

        helpmenu = Menu(topmenu, tearoff=0)
        helpmenu.add_command(label="About", command=self.aboutMenu)
        topmenu.add_cascade(label="Help", menu=helpmenu)

        # show menu
        self.tkapp.config(menu=topmenu)


    def openMenu(self):
        # set initial directory to savedScans folder
        path = os.path.join(os.getcwd(), "Sessions")
        if not os.path.exists(path):
            os.makedirs(path)
        filename = filedialog.askopenfilename(initialdir=path)
        if os.path.exists(path+filename):
            self.saveFilename = filename


    def saveMenu(self, saveType):
        # create main directory and subdir(current date) if not made already
        path = os.getcwd() + "/Sessions/" + str(datetime.date.today())
        if not os.path.exists(path):
            os.makedirs(path)

        if saveType == "Save" and self.saveFilename != None:
            if messagebox.askyesnocancel("Overwrite File?", "Overwrite {}?".format(self.saveFilename)):
                with open(os.path.join(path, self.saveFilename), 'w') as file:
                    file.write("Testing file already saved, resaving")
        else:
            # get a filename from the user or default to current time
            currentTime = datetime.datetime.now().strftime("%H_%M_%S")

            filename = filedialog.asksaveasfilename(defaultextension="txt", initialdir=path, initialfile=currentTime)
            if filename:
                self.saveFilename = filename
                with open(filename, 'w') as f:
                    f.write("Testing Save As/No Current Save")

    def exportMenu(self, filetype):
        keyword = 'Skype Fixes \u2018SPYKE\u2019 Credential Phishing Remote Execution Bug'
        url = 'http://cbrown686-test.apigee.net/cyberapi/articles?q=keywordtitle&title={}&author=&sub=&sdate=01/01/0001&edate=04/24/2017'.format(keyword)
        res = requests.get(url).json()
        res = res[0]
        if filetype == 'pdf':
            style = getSampleStyleSheet()
            sty = getSampleStyleSheet()['Title']
            sty.leading = 80
            pdf = SimpleDocTemplate("export.pdf", pagesize=letter)
            article = []

            article.append(Spacer(0,inch))
            article.append(Paragraph("<para fontSize=40 spaceAfter=40>{}</para>".format(res['title']), sty))
            article.append(Paragraph("<para fontSize=20 align=center>Author: {}</para>".format(res['author']), style['Normal']))
            article.append(Spacer(0,inch))
            article.append(Paragraph("<para fontSize=15 align=center spaceAfter=70>Date Published: {}</para>".format(res['date']), style["Normal"]))
            article.append(
                Paragraph("<para alignment=center><link fontSize=12 textColor=blue "
                          "href={}><u>{}</u></link></para>".format(res['uri'], res['uri']),style['Normal']))
            article.append(PageBreak())

            for para in res['body'].split('\n\n'):
                article.append(Paragraph(para, style['Normal']))
                article.append(Spacer(.25*inch, .25*inch))
            pdf.build(article)
        if filetype == 'docx':
            document = Document()
            document.add_heading(res['title'], 0)
            document.add_heading("Author: {}".format(res['author']), 1)
            document.add_heading("Date Published: {}".format(res['date']), 1)
            document.add_page_break()
            p = document.add_paragraph()
            for para in res['body'].split('\n\n'):
                document.add_paragraph(para.replace('\n', ''))
            document.save('export.docx')

    def aboutMenu(self):
        messagebox.showinfo("About", "Developed by: Chase Brown and Joseph Dodson\n"
                                     "Refer to README for additional info.")

